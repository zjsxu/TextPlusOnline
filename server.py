from flask import Flask, request, jsonify
from PyPDF2 import PdfReader 
from docx import Document # 【新增】用于读取 DOCX 文件
import io
import os 
from flask_cors import CORS 
import json

app = Flask(__name__)
# -------------------------
# 载入功能词词典（来自导师 / doccano）
# -------------------------
DICT_PATH = os.path.join(os.path.dirname(__file__), "dictionary/function_words.json")

def load_function_dict():
    try:
        with open(DICT_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        print("词典加载失败:", e)
        return {}

FUNCTION_WORD_DICT = load_function_dict()

CORS(app) 

# 定义文件文本提取的 API 路由
@app.route('/')
def index():
    return "Text-Diff-Tool 后端已运行。请使用前端界面上传文件。"

@app.route('/get-dict', methods=['GET'])
def get_dict():
    return jsonify(FUNCTION_WORD_DICT)

@app.route('/extract-text', methods=['POST'])
def extract_text():
    if 'file' not in request.files:
        return jsonify({'error': '未找到文件部分'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '未选择文件'}), 400

    filename = file.filename
    file_extension = os.path.splitext(filename)[1].lower()
    text_content = ""

    try:
        # --- TXT 文件处理 ---
        if file_extension == '.txt':
            text_content = file.read().decode('utf-8')

        # --- PDF 文件处理 ---
        elif file_extension == '.pdf':
            reader = PdfReader(io.BytesIO(file.read()))
            for page in reader.pages:
                extracted_page_text = page.extract_text()
                if extracted_page_text:
                     # 使用自定义分隔符，防止文本内容完全粘在一起
                     text_content += extracted_page_text + "\n\n--- 页面分割线 ---\n\n"
            
            if not text_content:
                return jsonify({'error': 'PDF 文件为空或无法提取文本，请确保内容不是纯图片'}), 400

        # --- 【新增】DOCX 文件处理 ---
        elif file_extension == '.docx':
            # 将文件内容读入内存流
            doc_stream = io.BytesIO(file.read())
            # 使用 python-docx 库读取文档
            document = Document(doc_stream)
            
            # 遍历文档中的所有段落并提取文本
            for paragraph in document.paragraphs:
                text_content += paragraph.text + "\n"
                
            # 处理表格内容 (可选，但推荐)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + "\t"
                text_content += "\n\n--- 表格分割线 ---\n\n"
            
            if not text_content:
                return jsonify({'error': 'DOCX 文件为空或无法提取文本'}), 400

        else:
            return jsonify({'error': f'不支持的文件类型: {file_extension}。仅支持 .pdf、.docx 和 .txt'}), 400

        return jsonify({'text': text_content})

    except Exception as e:
        print(f"提取过程中发生错误: {e}")
        return jsonify({'error': f'文件处理失败，请检查文件格式。具体错误: {str(e)}'}), 500

if __name__ == '__main__':
    print("Python 后端服务器已启动在 http://127.0.0.1:5000")
    app.run(debug=True, port=5000,use_reloader=False)